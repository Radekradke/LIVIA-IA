"""Geração de documentos e planilhas reais.

Até aqui a Livia escrevia texto puro. Aqui ela produz arquivos que abrem no
Word, no Excel e em qualquer leitor de PDF.

TRÊS GARANTIAS QUE VALEM MAIS QUE O RECURSO
-------------------------------------------
1. ESCRITA ATÔMICA. O arquivo é montado num temporário e só então renomeado
   para o destino. Se a geração falhar no meio, o que existia continua
   intacto e nenhum arquivo pela metade aparece na pasta.

2. FÓRMULA NÃO ENTRA EM PLANILHA. Texto começando com `=`, `+`, `-` ou `@`
   vira texto literal, nunca fórmula. Isso é defesa contra injeção de
   fórmula (CSV injection): uma célula com `=cmd|...` pode virar execução de
   comando quando alguém abre a planilha no Excel. O conteúdo vem do modelo;
   confiar seria irresponsável.

3. CONFINAMENTO. Mesmo `_resolver` das outras ferramentas — nada escapa da
   pasta de trabalho, e sobrescrever guarda cópia da versão anterior.

MARKDOWN, SÓ O ESSENCIAL
------------------------
Títulos, parágrafos e listas. Não é um conversor completo, e tentar ser um
seria construir um projeto dentro do projeto. O que não for reconhecido vira
parágrafo comum — degradação previsível em vez de erro.
"""

from __future__ import annotations

import csv
import html as _html
import re
from pathlib import Path

from .ferramentas import FerramentaError, _resolver, _tamanho, raiz

LIMITE_CONTEUDO = 300_000     # caracteres do texto de entrada
LIMITE_LINHAS = 5_000         # linhas por planilha
LIMITE_COLUNAS = 100

FORMATOS_DOC = ("txt", "md", "html", "pdf", "docx")
FORMATOS_PLANILHA = ("xlsx", "csv")

# Uma célula que começa com um destes é interpretada como fórmula pelo Excel
# e pelo Google Sheets. Prefixar com aspa simples neutraliza sem mudar o que
# a pessoa lê na tela.
_GATILHOS_FORMULA = ("=", "+", "-", "@", "\t", "\r")


# --------------------------------------------------------------------------
# Markdown mínimo
# --------------------------------------------------------------------------


def _blocos(texto: str) -> list[tuple[str, str, int]]:
    """Quebra o texto em (tipo, conteúdo, nível).

    Tipos: `titulo`, `lista`, `numerada`, `paragrafo`.
    """
    saida: list[tuple[str, str, int]] = []
    for linha in (texto or "").replace("\r\n", "\n").split("\n"):
        bruta = linha.rstrip()
        limpa = bruta.strip()
        if not limpa:
            continue

        cabecalho = re.match(r"^(#{1,6})\s+(.*)$", limpa)
        if cabecalho:
            saida.append(("titulo", _inline(cabecalho.group(2)), len(cabecalho.group(1))))
            continue

        marcador = re.match(r"^[-*+]\s+(.*)$", limpa)
        if marcador:
            saida.append(("lista", _inline(marcador.group(1)), 0))
            continue

        numerada = re.match(r"^\d+[.)]\s+(.*)$", limpa)
        if numerada:
            saida.append(("numerada", _inline(numerada.group(1)), 0))
            continue

        saida.append(("paragrafo", _inline(limpa), 0))
    return saida


def _inline(texto: str) -> str:
    """Tira a marcação de ênfase. Sem negrito por enquanto — melhor um texto
    limpo do que asteriscos soltos no meio do documento."""
    texto = re.sub(r"\*\*(.+?)\*\*", r"\1", texto)
    texto = re.sub(r"(?<!\*)\*(?!\*)(.+?)(?<!\*)\*(?!\*)", r"\1", texto)
    texto = re.sub(r"`(.+?)`", r"\1", texto)
    return texto


# --------------------------------------------------------------------------
# Escrita atômica
# --------------------------------------------------------------------------


def _preparar(caminho: str, formato: str, permitidos: tuple[str, ...]) -> tuple[Path, Path]:
    """Valida tudo e devolve (destino final, arquivo temporário)."""
    formato = (formato or "").strip().lower().lstrip(".")
    if formato not in permitidos:
        raise FerramentaError(
            f"Formato '{formato}' não serve aqui. Use: {', '.join(permitidos)}."
        )

    alvo = _resolver(caminho)
    extensao = alvo.suffix.lower().lstrip(".")
    if extensao != formato:
        raise FerramentaError(
            f"A extensão de '{caminho}' é '.{extensao}' mas o formato pedido é "
            f"'{formato}'. Faça os dois combinarem."
        )
    if alvo.is_dir():
        raise FerramentaError(f"'{caminho}' é uma pasta.")

    alvo.parent.mkdir(parents=True, exist_ok=True)
    temporario = alvo.with_name(f".{alvo.name}.parcial")
    return alvo, temporario


def _publicar(alvo: Path, temporario: Path, formato: str) -> str:
    """Move o temporário para o destino, guardando a versão anterior."""
    from datetime import datetime

    if not temporario.exists():
        raise FerramentaError("A geração não produziu arquivo nenhum.")

    if alvo.exists():
        carimbo = datetime.now().strftime("%Y%m%d-%H%M%S")
        copia = alvo.with_name(f"{alvo.stem}.antes-{carimbo}{alvo.suffix}")
        try:
            alvo.replace(copia)
        except OSError:
            pass

    temporario.replace(alvo)
    tamanho = alvo.stat().st_size
    relativo = alvo.relative_to(raiz()).as_posix()
    return f"Criei '{relativo}' ({formato.upper()}, {_tamanho(tamanho)})."


def _descartar(temporario: Path) -> None:
    """Some com o parcial. Arquivo pela metade confunde mais que ajuda."""
    try:
        if temporario.exists():
            temporario.unlink()
    except OSError:
        pass


# --------------------------------------------------------------------------
# Documentos
# --------------------------------------------------------------------------


def criar_documento(caminho: str, titulo: str, conteudo: str, formato: str = "md") -> str:
    if len(conteudo or "") > LIMITE_CONTEUDO:
        raise FerramentaError(
            f"Conteúdo grande demais ({len(conteudo)} caracteres). "
            f"O teto é {LIMITE_CONTEUDO}."
        )

    alvo, temporario = _preparar(caminho, formato, FORMATOS_DOC)
    formato = alvo.suffix.lower().lstrip(".")
    titulo = (titulo or "").strip()

    try:
        if formato in ("txt", "md"):
            _escrever_texto(temporario, titulo, conteudo, formato)
        elif formato == "html":
            _escrever_html(temporario, titulo, conteudo)
        elif formato == "docx":
            _escrever_docx(temporario, titulo, conteudo)
        elif formato == "pdf":
            _escrever_pdf(temporario, titulo, conteudo)
        return _publicar(alvo, temporario, formato)
    except FerramentaError:
        _descartar(temporario)
        raise
    except Exception as exc:
        _descartar(temporario)
        raise FerramentaError(f"Não consegui gerar o {formato.upper()}: {exc}") from exc


def _escrever_texto(destino: Path, titulo: str, conteudo: str, formato: str) -> None:
    partes = []
    if titulo:
        partes.append(f"# {titulo}\n" if formato == "md" else f"{titulo}\n{'=' * len(titulo)}\n")
    partes.append(conteudo.strip() + "\n")
    destino.write_text("\n".join(partes), encoding="utf-8")


def _escrever_html(destino: Path, titulo: str, conteudo: str) -> None:
    corpo: list[str] = []
    lista_aberta: str | None = None

    def fechar():
        nonlocal lista_aberta
        if lista_aberta:
            corpo.append(f"</{lista_aberta}>")
            lista_aberta = None

    for tipo, texto, nivel in _blocos(conteudo):
        seguro = _html.escape(texto)
        if tipo == "titulo":
            fechar()
            corpo.append(f"<h{min(nivel + 1, 6)}>{seguro}</h{min(nivel + 1, 6)}>")
        elif tipo in ("lista", "numerada"):
            marca = "ul" if tipo == "lista" else "ol"
            if lista_aberta != marca:
                fechar()
                corpo.append(f"<{marca}>")
                lista_aberta = marca
            corpo.append(f"<li>{seguro}</li>")
        else:
            fechar()
            corpo.append(f"<p>{seguro}</p>")
    fechar()

    destino.write_text(
        "<!doctype html>\n<html lang=\"pt-BR\">\n<head>\n"
        '<meta charset="utf-8">\n'
        f"<title>{_html.escape(titulo or destino.stem)}</title>\n"
        "<style>body{font:16px/1.7 system-ui,sans-serif;max-width:46em;"
        "margin:3rem auto;padding:0 1.5rem;color:#222}"
        "h1,h2,h3{line-height:1.3}code{background:#f2f2f2;padding:.1em .3em}</style>\n"
        "</head>\n<body>\n"
        + (f"<h1>{_html.escape(titulo)}</h1>\n" if titulo else "")
        + "\n".join(corpo)
        + "\n</body>\n</html>\n",
        encoding="utf-8",
    )


def _escrever_docx(destino: Path, titulo: str, conteudo: str) -> None:
    from docx import Document

    doc = Document()
    if titulo:
        doc.add_heading(titulo, level=0)

    for tipo, texto, nivel in _blocos(conteudo):
        if tipo == "titulo":
            doc.add_heading(texto, level=min(nivel, 4))
        elif tipo == "lista":
            doc.add_paragraph(texto, style="List Bullet")
        elif tipo == "numerada":
            doc.add_paragraph(texto, style="List Number")
        else:
            doc.add_paragraph(texto)

    doc.save(str(destino))


def _escrever_pdf(destino: Path, titulo: str, conteudo: str) -> None:
    from reportlab.lib.enums import TA_LEFT
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
    from reportlab.lib.units import cm
    from reportlab.platypus import ListFlowable, ListItem, Paragraph, SimpleDocTemplate, Spacer

    estilos = getSampleStyleSheet()
    corpo = ParagraphStyle(
        "corpo", parent=estilos["BodyText"], fontSize=11, leading=16,
        spaceAfter=8, alignment=TA_LEFT,
    )

    doc = SimpleDocTemplate(
        str(destino), pagesize=A4,
        leftMargin=2.5 * cm, rightMargin=2.5 * cm,
        topMargin=2.5 * cm, bottomMargin=2.5 * cm,
        title=titulo or destino.stem,
    )

    fluxo = []
    if titulo:
        fluxo += [Paragraph(_html.escape(titulo), estilos["Title"]), Spacer(1, 8)]

    pendentes: list = []
    marca_atual: str | None = None

    def fechar_lista():
        nonlocal pendentes, marca_atual
        if pendentes:
            fluxo.append(
                ListFlowable(
                    pendentes,
                    bulletType="bullet" if marca_atual == "lista" else "1",
                    leftIndent=18,
                )
            )
            fluxo.append(Spacer(1, 6))
            pendentes = []
            marca_atual = None

    for tipo, texto, nivel in _blocos(conteudo):
        seguro = _html.escape(texto)
        if tipo == "titulo":
            fechar_lista()
            estilo = estilos["Heading1"] if nivel <= 1 else estilos[f"Heading{min(nivel, 4)}"]
            fluxo.append(Paragraph(seguro, estilo))
        elif tipo in ("lista", "numerada"):
            if marca_atual and marca_atual != tipo:
                fechar_lista()
            marca_atual = tipo
            pendentes.append(ListItem(Paragraph(seguro, corpo), leftIndent=18))
        else:
            fechar_lista()
            fluxo.append(Paragraph(seguro, corpo))
    fechar_lista()

    if not fluxo:
        fluxo.append(Paragraph("(documento vazio)", corpo))
    doc.build(fluxo)


# --------------------------------------------------------------------------
# Planilhas
# --------------------------------------------------------------------------


def _neutralizar(valor: object) -> object:
    """Impede que uma célula de texto vire fórmula ao abrir no Excel.

    `=cmd|'/c calc'!A1` numa célula é execução de comando quando alguém abre
    o arquivo. O conteúdo vem do modelo — confiar seria irresponsável. A aspa
    simples na frente neutraliza e não aparece para quem lê.
    """
    if isinstance(valor, (int, float, bool)) or valor is None:
        return valor
    texto = str(valor)
    if texto.startswith(_GATILHOS_FORMULA):
        return "'" + texto
    return texto


def criar_planilha(
    caminho: str,
    titulo: str,
    colunas: list[object],
    linhas: list[list[object]],
    formato: str = "xlsx",
) -> str:
    colunas = list(colunas or [])
    linhas = [list(l) for l in (linhas or [])]

    if not colunas:
        raise FerramentaError("A planilha precisa de pelo menos uma coluna.")
    if len(colunas) > LIMITE_COLUNAS:
        raise FerramentaError(f"Colunas demais ({len(colunas)}); o teto é {LIMITE_COLUNAS}.")
    if len(linhas) > LIMITE_LINHAS:
        raise FerramentaError(f"Linhas demais ({len(linhas)}); o teto é {LIMITE_LINHAS}.")

    for i, linha in enumerate(linhas, start=1):
        if len(linha) > len(colunas):
            raise FerramentaError(
                f"A linha {i} tem {len(linha)} valores para {len(colunas)} colunas."
            )

    alvo, temporario = _preparar(caminho, formato, FORMATOS_PLANILHA)
    formato = alvo.suffix.lower().lstrip(".")

    try:
        if formato == "csv":
            _escrever_csv(temporario, colunas, linhas)
        else:
            _escrever_xlsx(temporario, titulo, colunas, linhas)
        return _publicar(alvo, temporario, formato)
    except FerramentaError:
        _descartar(temporario)
        raise
    except Exception as exc:
        _descartar(temporario)
        raise FerramentaError(f"Não consegui gerar a planilha: {exc}") from exc


def _escrever_csv(destino: Path, colunas: list, linhas: list[list]) -> None:
    # utf-8-sig: sem o BOM o Excel abre acentuação errada no Windows.
    with destino.open("w", encoding="utf-8-sig", newline="") as f:
        escritor = csv.writer(f, delimiter=";")
        escritor.writerow([_neutralizar(c) for c in colunas])
        for linha in linhas:
            completa = list(linha) + [""] * (len(colunas) - len(linha))
            escritor.writerow([_neutralizar(v) for v in completa])


def _escrever_xlsx(destino: Path, titulo: str, colunas: list, linhas: list[list]) -> None:
    from openpyxl import Workbook
    from openpyxl.styles import Alignment, Font, PatternFill
    from openpyxl.utils import get_column_letter

    wb = Workbook()
    aba = wb.active
    # O Excel recusa alguns caracteres em nome de aba, e trunca em 31.
    aba.title = re.sub(r"[\\/*?\[\]:]", "-", (titulo or "Dados").strip())[:31] or "Dados"

    aba.append([_neutralizar(c) for c in colunas])
    for linha in linhas:
        completa = list(linha) + [None] * (len(colunas) - len(linha))
        aba.append([_neutralizar(v) for v in completa])

    negrito = Font(bold=True, color="FFFFFF")
    fundo = PatternFill("solid", fgColor="8E0F2A")
    for celula in aba[1]:
        celula.font = negrito
        celula.fill = fundo
        celula.alignment = Alignment(vertical="center")

    for indice, coluna in enumerate(colunas, start=1):
        maior = len(str(coluna))
        for linha in linhas:
            if indice <= len(linha):
                maior = max(maior, len(str(linha[indice - 1])))
        aba.column_dimensions[get_column_letter(indice)].width = min(maior + 4, 60)

    aba.freeze_panes = "A2"
    wb.save(str(destino))


# --------------------------------------------------------------------------
# Catálogo para o modelo
# --------------------------------------------------------------------------

CATALOGO: list[dict[str, object]] = [
    {
        "name": "criar_documento",
        "description": (
            "Cria um documento de verdade na pasta de trabalho: TXT, Markdown, "
            "HTML, PDF ou DOCX. Use quando o André pedir um relatório, resumo, "
            "carta ou qualquer entregável para abrir em outro programa. "
            "A extensão do caminho tem que combinar com o formato."
        ),
        "parameters": {
            "type": "object",
            "properties": {
                "caminho": {
                    "type": "string",
                    "description": "Relativo à pasta de trabalho, ex: relatorios/analise.docx",
                },
                "titulo": {"type": "string", "description": "Título que abre o documento."},
                "conteudo": {
                    "type": "string",
                    "description": (
                        "Corpo em Markdown simples: '# título', '- item', "
                        "'1. item' e parágrafos. Escreva o texto completo."
                    ),
                },
                "formato": {
                    "type": "string",
                    "enum": list(FORMATOS_DOC),
                    "description": "txt, md, html, pdf ou docx.",
                },
            },
            "required": ["caminho", "conteudo", "formato"],
        },
    },
    {
        "name": "criar_planilha",
        "description": (
            "Cria uma planilha XLSX ou CSV na pasta de trabalho. Use para "
            "tabelas, orçamentos, listas com colunas. Fórmulas não são "
            "aceitas: todo texto entra como texto."
        ),
        "parameters": {
            "type": "object",
            "properties": {
                "caminho": {
                    "type": "string",
                    "description": "Relativo à pasta de trabalho, ex: relatorios/custos.xlsx",
                },
                "titulo": {"type": "string", "description": "Nome da aba."},
                "colunas": {
                    "type": "array",
                    "items": {"type": "string"},
                    "description": "Cabeçalhos, na ordem.",
                },
                "linhas": {
                    "type": "array",
                    "items": {"type": "array", "items": {"type": "string"}},
                    "description": "Cada item é uma linha, com valores na ordem das colunas.",
                },
                "formato": {
                    "type": "string",
                    "enum": list(FORMATOS_PLANILHA),
                    "description": "xlsx ou csv.",
                },
            },
            "required": ["caminho", "colunas", "linhas", "formato"],
        },
    },
]

EXECUTORES = {
    "criar_documento": lambda a: criar_documento(
        str(a.get("caminho") or ""),
        str(a.get("titulo") or ""),
        str(a.get("conteudo") or ""),
        str(a.get("formato") or "md"),
    ),
    "criar_planilha": lambda a: criar_planilha(
        str(a.get("caminho") or ""),
        str(a.get("titulo") or ""),
        list(a.get("colunas") or []),
        list(a.get("linhas") or []),
        str(a.get("formato") or "xlsx"),
    ),
}


def resumir(nome: str, argumentos: dict[str, object]) -> str:
    caminho = argumentos.get("caminho") or "arquivo"
    if nome == "criar_documento":
        return f"criando documento {caminho}"
    if nome == "criar_planilha":
        n = len(argumentos.get("linhas") or [])
        return f"criando planilha {caminho} ({n} linhas)"
    return f"usando {nome}"
