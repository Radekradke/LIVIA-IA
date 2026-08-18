"""Extração de texto por formato — o par que faltava para a fase 5.

Ela já escrevia DOCX; agora lê o que o André jogar na pasta. PDF, DOCX,
XLSX, CSV, JSON e HTML entram como texto legível para o modelo.

DUAS COISAS QUE UM LEITOR INGÊNUO ERRA
--------------------------------------
1. XLSX abre com `data_only=True`. Sem isso o openpyxl devolve a FÓRMULA
   (`=SOMA(A1:A9)`) em vez do resultado, e a Livia responderia comentando o
   texto da fórmula achando que é o dado. Também importa que a leitura nunca
   calcula nada: só lê o valor que o Excel gravou.

2. DOCX e XLSX são arquivos ZIP. Um `.docx` de 50 KB pode virar 4 GB ao
   descompactar e derrubar o processo — isso é bomba-zip, e existe antes de
   qualquer má intenção do modelo, porque o arquivo vem de fora. Conferimos
   o tamanho declarado ANTES de extrair.

O QUE NÃO ESTÁ AQUI
-------------------
OCR. Um PDF escaneado não tem camada de texto, e nenhuma biblioteca pura de
Python resolve isso. Em vez de devolver vazio e deixar a Livia inventar o
conteúdo, dizemos exatamente o que aconteceu.
"""

from __future__ import annotations

import csv
import html as _html
import io
import json
import re
import zipfile
from pathlib import Path

from .ferramentas import FerramentaError, _tamanho

LIMITE_ARQUIVO = 25 * 1024 * 1024      # 25 MB de arquivo em disco
LIMITE_DESCOMPACTADO = 80 * 1024 * 1024  # o que um DOCX/XLSX pode virar
RAZAO_SUSPEITA = 150                    # descompactado ÷ comprimido
LIMITE_LINHAS_TABELA = 400              # linhas por aba/tabela mostradas
LIMITE_PAGINAS = 200

# Extensões com tratamento próprio. Fora desta lista, tentamos ler como texto.
EXTENSOES = ("pdf", "docx", "xlsx", "xlsm", "csv", "tsv", "json", "html", "htm")


def suportado(caminho: Path) -> bool:
    return caminho.suffix.lower().lstrip(".") in EXTENSOES


def extrair(caminho: Path, rotulo: str) -> str:
    """Texto legível de um arquivo estruturado. `rotulo` é o nome que o
    usuário usou, para a mensagem de erro falar a língua dele."""
    tamanho = caminho.stat().st_size
    if tamanho > LIMITE_ARQUIVO:
        raise FerramentaError(
            f"'{rotulo}' tem {_tamanho(tamanho)} — grande demais para eu ler "
            f"inteiro (o teto é {_tamanho(LIMITE_ARQUIVO)})."
        )

    extensao = caminho.suffix.lower().lstrip(".")
    if extensao == "pdf":
        return _pdf(caminho, rotulo)
    if extensao == "docx":
        return _docx(caminho, rotulo)
    if extensao in ("xlsx", "xlsm"):
        return _xlsx(caminho, rotulo)
    if extensao in ("csv", "tsv"):
        return _csv(caminho, rotulo)
    if extensao == "json":
        return _json(caminho, rotulo)
    if extensao in ("html", "htm"):
        return _html_texto(caminho)
    raise FerramentaError(f"Não sei ler arquivos '.{extensao}'.")


# --------------------------------------------------------------------------
# Bomba-zip
# --------------------------------------------------------------------------


def conferir_zip(origem, rotulo: str) -> None:
    """Recusa antes de descompactar. O arquivo veio de fora.

    `origem` é o que o zipfile aceita: um caminho ou bytes em memória. A
    biblioteca recebe upload, que nunca toca o disco antes da conferência.
    """
    try:
        with zipfile.ZipFile(origem) as z:
            itens = z.infolist()
            descompactado = sum(i.file_size for i in itens)
            compactado = sum(i.compress_size for i in itens) or 1
    except zipfile.BadZipFile:
        raise FerramentaError(
            f"'{rotulo}' não é um arquivo válido desse tipo — DOCX e XLSX são "
            "pacotes ZIP por dentro, e este está corrompido ou tem outro formato."
        ) from None

    if descompactado > LIMITE_DESCOMPACTADO:
        raise FerramentaError(
            f"'{rotulo}' declara {_tamanho(descompactado)} descompactado. "
            "Não vou abrir."
        )
    if descompactado / compactado > RAZAO_SUSPEITA:
        raise FerramentaError(
            f"'{rotulo}' expande {descompactado // compactado}x ao descompactar. "
            "Isso é característico de bomba-zip; não vou abrir."
        )


# --------------------------------------------------------------------------
# PDF
# --------------------------------------------------------------------------


def _pdf(caminho: Path, rotulo: str) -> str:
    from pypdf import PdfReader
    from pypdf.errors import PdfReadError

    try:
        leitor = PdfReader(str(caminho))
        if leitor.is_encrypted:
            try:
                leitor.decrypt("")  # muitos PDFs só têm senha de dono
            except Exception:
                pass
            if leitor.is_encrypted:
                raise FerramentaError(
                    f"'{rotulo}' está protegido por senha. Remova a proteção "
                    "e mande de novo."
                )
        paginas = leitor.pages
    except FerramentaError:
        raise
    except (PdfReadError, Exception) as exc:
        raise FerramentaError(f"Não consegui abrir '{rotulo}' como PDF: {exc}") from exc

    partes: list[str] = []
    total = len(paginas)
    for numero, pagina in enumerate(paginas[:LIMITE_PAGINAS], start=1):
        try:
            texto = (pagina.extract_text() or "").strip()
        except Exception:
            texto = ""
        if texto:
            partes.append(f"--- página {numero} ---\n{texto}")

    if not partes:
        # Devolver vazio faria a Livia comentar um documento que ela não leu.
        raise FerramentaError(
            f"'{rotulo}' tem {total} página(s) mas nenhum texto extraível — "
            "provavelmente é um escaneamento (imagem). Eu não faço OCR, então "
            "não consigo ler o conteúdo. Se tiver a versão original em texto, "
            "ela funciona."
        )

    cabecalho = f"[PDF: {total} página(s)]"
    if total > LIMITE_PAGINAS:
        cabecalho += f" (li as primeiras {LIMITE_PAGINAS})"
    return cabecalho + "\n\n" + "\n\n".join(partes)


# --------------------------------------------------------------------------
# DOCX
# --------------------------------------------------------------------------


def _docx(caminho: Path, rotulo: str) -> str:
    conferir_zip(caminho, rotulo)
    from docx import Document

    try:
        doc = Document(str(caminho))
    except Exception as exc:
        raise FerramentaError(f"Não consegui abrir '{rotulo}' como DOCX: {exc}") from exc

    partes: list[str] = []
    for paragrafo in doc.paragraphs:
        texto = paragrafo.text.strip()
        if not texto:
            continue
        estilo = (paragrafo.style.name or "").lower()
        if estilo.startswith(("heading", "título", "titulo")):
            nivel = "".join(c for c in estilo if c.isdigit())
            partes.append(f"{'#' * (int(nivel) if nivel else 1)} {texto}")
        elif "list" in estilo:
            partes.append(f"- {texto}")
        else:
            partes.append(texto)

    # A tabela costuma ser onde o dado está. Ignorá-la seria perder o miolo.
    for indice, tabela in enumerate(doc.tables, start=1):
        linhas = []
        for linha in tabela.rows[:LIMITE_LINHAS_TABELA]:
            linhas.append(" | ".join(c.text.strip().replace("\n", " ") for c in linha.cells))
        if linhas:
            extra = "" if len(tabela.rows) <= LIMITE_LINHAS_TABELA else " (truncada)"
            partes.append(f"\n[tabela {indice}{extra}]\n" + "\n".join(linhas))

    if not partes:
        return "(o documento está vazio)"

    # Item de lista cola no anterior; o resto ganha linha em branco. Separar
    # tudo com parágrafo vazio dobraria o tamanho de um documento longo — e
    # isso é janela de contexto gasta à toa.
    corpo = partes[0]
    for parte in partes[1:]:
        anterior = corpo.rsplit("\n", 1)[-1]
        junta = "\n" if parte.startswith("- ") and anterior.startswith("- ") else "\n\n"
        corpo += junta + parte

    return (
        f"[DOCX: {len(doc.paragraphs)} parágrafos, {len(doc.tables)} tabela(s)]\n\n"
        + corpo
    )


# --------------------------------------------------------------------------
# XLSX
# --------------------------------------------------------------------------


def _xlsx(caminho: Path, rotulo: str) -> str:
    conferir_zip(caminho, rotulo)
    from openpyxl import load_workbook

    try:
        # data_only=True: queremos o RESULTADO gravado, não o texto da fórmula.
        # read_only=True: não carrega a planilha inteira na memória de uma vez.
        wb = load_workbook(str(caminho), data_only=True, read_only=True)
    except Exception as exc:
        raise FerramentaError(f"Não consegui abrir '{rotulo}' como planilha: {exc}") from exc

    partes: list[str] = []
    try:
        for aba in wb.worksheets:
            linhas: list[str] = []
            vazia = True
            for indice, linha in enumerate(aba.iter_rows(values_only=True), start=1):
                if indice > LIMITE_LINHAS_TABELA:
                    linhas.append(f"[... aba truncada em {LIMITE_LINHAS_TABELA} linhas]")
                    break
                celulas = ["" if c is None else str(c) for c in linha]
                if any(celulas):
                    vazia = False
                linhas.append(" | ".join(celulas).rstrip(" |"))

            while linhas and not linhas[-1].strip():
                linhas.pop()

            titulo = f"--- aba: {aba.title} ---"
            if vazia:
                # Planilha gerada por programa e nunca aberta no Excel não tem
                # resultado de fórmula gravado. Dizer isso evita a Livia
                # concluir que a planilha está vazia.
                partes.append(
                    titulo + "\n(sem valores legíveis — se a aba usa fórmulas e o "
                    "arquivo nunca foi aberto no Excel, os resultados não ficam "
                    "gravados no arquivo)"
                )
            else:
                partes.append(titulo + "\n" + "\n".join(linhas))
    finally:
        wb.close()

    if not partes:
        return "(a planilha não tem abas)"
    return f"[Planilha: {len(partes)} aba(s)]\n\n" + "\n\n".join(partes)


# --------------------------------------------------------------------------
# CSV
# --------------------------------------------------------------------------


def _csv(caminho: Path, rotulo: str) -> str:
    bruto = _texto_bruto(caminho, rotulo)
    if not bruto.strip():
        return "(o arquivo está vazio)"

    amostra = bruto[:8000]
    try:
        separador = csv.Sniffer().sniff(amostra, delimiters=";,\t|").delimiter
    except csv.Error:
        # No Brasil o Excel usa `;`. Escolher pela contagem erra menos que
        # assumir a vírgula do padrão americano.
        separador = max(";,\t|", key=lambda d: amostra.count(d))

    linhas = list(csv.reader(io.StringIO(bruto), delimiter=separador))
    linhas = [l for l in linhas if any(str(c).strip() for c in l)]
    if not linhas:
        return "(o arquivo está vazio)"

    total = len(linhas)
    mostradas = linhas[:LIMITE_LINHAS_TABELA]
    corpo = "\n".join(" | ".join(str(c) for c in l) for l in mostradas)

    cabecalho = f"[CSV: {total} linha(s), {len(linhas[0])} coluna(s), separador '{separador}']"
    if total > LIMITE_LINHAS_TABELA:
        corpo += f"\n[... e mais {total - LIMITE_LINHAS_TABELA} linhas]"
    return cabecalho + "\n\n" + corpo


# --------------------------------------------------------------------------
# JSON
# --------------------------------------------------------------------------


def _json(caminho: Path, rotulo: str) -> str:
    bruto = _texto_bruto(caminho, rotulo)
    try:
        dados = json.loads(bruto)
    except json.JSONDecodeError as exc:
        # Devolver o texto cru ainda serve: o modelo costuma achar o erro.
        return (
            f"[JSON inválido: {exc.msg} na linha {exc.lineno}, coluna {exc.colno}]\n\n"
            + bruto[:20_000]
        )

    formatado = json.dumps(dados, ensure_ascii=False, indent=2)
    if isinstance(dados, list):
        resumo = f"[JSON: lista com {len(dados)} itens]"
    elif isinstance(dados, dict):
        resumo = f"[JSON: objeto com as chaves {', '.join(list(dados)[:15])}]"
    else:
        resumo = "[JSON]"
    return resumo + "\n\n" + formatado


# --------------------------------------------------------------------------
# HTML
# --------------------------------------------------------------------------


def _html_texto(caminho: Path) -> str:
    bruto = _texto_bruto(caminho, caminho.name)

    titulo = ""
    achado = re.search(r"<title[^>]*>(.*?)</title>", bruto, re.I | re.S)
    if achado:
        titulo = _html.unescape(achado.group(1)).strip()

    # O conteúdo de script e style é código, não leitura. Tirar só as tags
    # deixaria o JavaScript inteiro no meio do texto.
    texto = re.sub(r"<(script|style)[^>]*>.*?</\1>", " ", bruto, flags=re.I | re.S)
    texto = re.sub(r"<!--.*?-->", " ", texto, flags=re.S)
    texto = re.sub(r"</(p|div|li|tr|h[1-6]|br)\s*>", "\n", texto, flags=re.I)
    texto = re.sub(r"<br\s*/?>", "\n", texto, flags=re.I)
    texto = re.sub(r"<[^>]+>", " ", texto)
    texto = _html.unescape(texto)
    texto = re.sub(r"[ \t]+", " ", texto)
    texto = re.sub(r"\n\s*\n\s*\n+", "\n\n", texto).strip()

    if not texto:
        return "(a página não tem texto)"
    return (f"[HTML: {titulo}]\n\n" if titulo else "[HTML]\n\n") + texto


# --------------------------------------------------------------------------


def _texto_bruto(caminho: Path, rotulo: str) -> str:
    dados = caminho.read_bytes()
    for codec in ("utf-8-sig", "utf-8", "cp1252", "latin-1"):
        try:
            return dados.decode(codec)
        except UnicodeDecodeError:
            continue
    raise FerramentaError(f"Não consegui decodificar '{rotulo}': o conteúdo não é texto.")
