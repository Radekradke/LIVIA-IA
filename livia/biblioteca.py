"""Biblioteca: livros e documentos que a Livia consulta ao responder.

É a mesma ideia da memória, um nível acima. O modelo não decora o livro — ele
é reiniciado a cada pergunta, como sempre. O que acontece é:

    ao adicionar   livro -> texto -> trechos -> vetores (guardados em disco)
    ao perguntar   pergunta -> vetor -> trechos mais parecidos -> vão no prompt

O "mais parecidos" é comparação de significado, não de palavras. Perguntar
"como reservo memória?" encontra o trecho sobre `malloc` mesmo sem a palavra
malloc aparecer na pergunta. É isso que faz parecer que ela leu o livro.

DECISÕES QUE VALEM SABER
------------------------
Vetores de 768 dimensões, não 3072. O modelo entrega 3072, mas ele é treinado
para que truncar mantenha o essencial. 768 ocupa um quarto do espaço e a perda
de qualidade é irrelevante nesta escala.

Busca sempre, injeta só se for relevante. Não perguntamos ao modelo "essa
pergunta é sobre algum livro?" — isso custaria uma chamada extra. Buscamos
sempre (barato) e só colamos no prompt o que passar de um limiar de
semelhança. Pergunta sem relação com os livros não injeta nada.
"""

from __future__ import annotations

import hashlib
import json
import re
import shutil
from collections.abc import AsyncIterator
from datetime import date
from pathlib import Path

import httpx
import numpy as np

from . import config
from .docs import slugify


def pasta_livros() -> Path:
    """Onde os livros ficam.

    Função, não constante: `config.DATA_DIR` pode ser trocado depois do
    import — nos testes é, e uma constante ignoraria a troca em silêncio,
    fazendo o teste gravar na pasta de dados real. A `raiz()` das
    ferramentas já funciona assim.
    """
    return config.DATA_DIR / "biblioteca"
EMBED_URL = "https://generativelanguage.googleapis.com/v1beta/models"

MODELO_EMBED = "gemini-embedding-001"
DIMENSOES = 768
TAMANHO_TRECHO = 1200      # caracteres
SOBREPOSICAO = 200         # evita cortar uma explicação ao meio
LOTE = 50                  # trechos por chamada à API
LIMIAR = 0.55              # abaixo disso, o trecho não tem a ver com a pergunta

EXTENSOES = {".pdf", ".docx", ".txt", ".md", ".markdown"}
LIMITE_ARQUIVO = 40 * 1024 * 1024


class BibliotecaError(RuntimeError):
    """Erro já em português, para mostrar ao usuário."""


# --------------------------------------------------------------------------
# Extrair texto
# --------------------------------------------------------------------------


def extrair(nome_arquivo: str, dados: bytes) -> list[tuple[int, str]]:
    """Devolve [(página, texto), ...]. Página é 0 em arquivos sem paginação."""
    ext = Path(nome_arquivo).suffix.lower()

    if ext == ".pdf":
        return _extrair_pdf(dados)
    if ext == ".docx":
        return _extrair_docx(dados)
    if ext in EXTENSOES:
        try:
            texto = dados.decode("utf-8")
        except UnicodeDecodeError:
            texto = dados.decode("latin-1", "replace")
        return [(0, texto)]

    raise BibliotecaError(
        f"Não sei ler '{ext}'. Aceito PDF, DOCX, TXT e Markdown. "
        "Para EPUB, converta antes."
    )


def _extrair_docx(dados: bytes) -> list[tuple[int, str]]:
    """DOCX não tem página — o Word pagina na hora de imprimir. Numeramos
    por seção, para a citação apontar algum lugar em vez de nenhum."""
    import io

    from docx import Document

    from . import leitura

    pacote = io.BytesIO(dados)
    try:
        # conferir_zip fala a língua das ferramentas; aqui a mensagem tem
        # que ser BibliotecaError, senão o servidor a mostra como
        # "Falha inesperada" e o motivo real se perde.
        leitura.conferir_zip(pacote, "o arquivo enviado")
    except Exception as exc:
        raise BibliotecaError(str(exc)) from exc
    pacote.seek(0)

    try:
        doc = Document(pacote)
    except Exception as exc:
        raise BibliotecaError(f"Não consegui abrir o DOCX: {exc}") from exc

    # Cada título abre uma seção nova. Isso dá à citação uma âncora que a
    # pessoa consegue achar no documento, mesmo sem número de página.
    secoes: list[tuple[int, str]] = []
    atual: list[str] = []
    numero = 1

    for paragrafo in doc.paragraphs:
        texto = paragrafo.text.strip()
        if not texto:
            continue
        estilo = (paragrafo.style.name or "").lower()
        if estilo.startswith(("heading", "título", "titulo")) and atual:
            secoes.append((numero, "\n\n".join(atual)))
            numero += 1
            atual = []
        atual.append(texto)

    for tabela in doc.tables:
        linhas = [
            " | ".join(c.text.strip().replace("\n", " ") for c in linha.cells)
            for linha in tabela.rows
        ]
        if any(l.strip(" |") for l in linhas):
            atual.append("\n".join(linhas))

    if atual:
        secoes.append((numero, "\n\n".join(atual)))

    if not secoes:
        raise BibliotecaError("Esse DOCX não tem texto nenhum.")
    return secoes


def _extrair_pdf(dados: bytes) -> list[tuple[int, str]]:
    import io

    from pypdf import PdfReader

    try:
        leitor = PdfReader(io.BytesIO(dados))
    except Exception as exc:
        raise BibliotecaError(f"Não consegui abrir o PDF: {exc}") from exc

    if leitor.is_encrypted:
        try:
            leitor.decrypt("")
        except Exception:
            raise BibliotecaError(
                "O PDF está protegido por senha. Remova a proteção e envie de novo."
            ) from None

    paginas: list[tuple[int, str]] = []
    for i, pagina in enumerate(leitor.pages, start=1):
        try:
            texto = pagina.extract_text() or ""
        except Exception:
            texto = ""
        if texto.strip():
            paginas.append((i, texto))

    if not paginas:
        raise BibliotecaError(
            "Não saiu texto nenhum desse PDF. Provavelmente é um livro "
            "escaneado (imagens de páginas, não texto). Esses precisam passar "
            "por OCR antes, e isso eu ainda não faço."
        )
    return paginas


# --------------------------------------------------------------------------
# Dividir em trechos
# --------------------------------------------------------------------------


def _limpar(texto: str) -> str:
    texto = texto.replace("\r\n", "\n").replace("\xa0", " ")
    texto = re.sub(r"[ \t]+", " ", texto)
    texto = re.sub(r"\n{3,}", "\n\n", texto)
    return texto.strip()


def dividir(paginas: list[tuple[int, str]]) -> list[dict[str, object]]:
    """Quebra o texto em pedaços que caibam no prompt, sem cortar no meio.

    A sobreposição existe porque uma explicação boa raramente respeita a
    fronteira que a gente inventou: sem ela, o trecho pode terminar em "e
    então você deve" e a resposta ficar pela metade.
    """
    trechos: list[dict[str, object]] = []

    for numero, bruto in paginas:
        texto = _limpar(bruto)
        if not texto:
            continue

        paragrafos = [p.strip() for p in texto.split("\n\n") if p.strip()]
        atual = ""

        for paragrafo in paragrafos:
            if len(atual) + len(paragrafo) + 2 <= TAMANHO_TRECHO:
                atual = f"{atual}\n\n{paragrafo}" if atual else paragrafo
                continue

            if atual:
                trechos.append({"pagina": numero, "texto": atual})
                atual = atual[-SOBREPOSICAO:] + "\n\n" + paragrafo
            else:
                atual = paragrafo

            # Parágrafo gigante (código, tabela): corta no tamanho mesmo.
            while len(atual) > TAMANHO_TRECHO:
                trechos.append({"pagina": numero, "texto": atual[:TAMANHO_TRECHO]})
                atual = atual[TAMANHO_TRECHO - SOBREPOSICAO :]

        if atual.strip():
            trechos.append({"pagina": numero, "texto": atual.strip()})

    # Trecho curto demais quase nunca responde nada.
    return [t for t in trechos if len(str(t["texto"])) > 120]


# --------------------------------------------------------------------------
# Vetores
# --------------------------------------------------------------------------


async def _embutir(textos: list[str], tarefa: str) -> list[list[float]]:
    """Converte textos em vetores. `tarefa` melhora a qualidade da busca."""
    if not config.GEMINI_API_KEY:
        raise BibliotecaError(
            "A biblioteca precisa da chave do Gemini (GEMINI_API_KEY no .env). "
            "A Groq não oferece geração de vetores."
        )

    modelo = MODELO_EMBED
    corpo = {
        "requests": [
            {
                "model": f"models/{modelo}",
                "content": {"parts": [{"text": t[:8000]}]},
                "taskType": tarefa,
                "outputDimensionality": DIMENSOES,
            }
            for t in textos
        ]
    }
    headers = {"x-goog-api-key": config.GEMINI_API_KEY, "Content-Type": "application/json"}

    async with httpx.AsyncClient(timeout=httpx.Timeout(180.0)) as client:
        resposta = await client.post(
            f"{EMBED_URL}/{modelo}:batchEmbedContents", json=corpo, headers=headers
        )
        if resposta.status_code == 429:
            raise BibliotecaError(
                "Bateu o limite da API no meio do processamento. Ele reseta "
                "sozinho — espere alguns minutos e envie o arquivo de novo."
            )
        if resposta.status_code >= 400:
            detalhe = resposta.json().get("error", {}).get("message", "")[:200]
            raise BibliotecaError(f"A API recusou o pedido: {detalhe}")
        return [e["values"] for e in resposta.json()["embeddings"]]


def _normalizar(m: np.ndarray) -> np.ndarray:
    """Deixa todo vetor com comprimento 1, para comparar por produto escalar."""
    normas = np.linalg.norm(m, axis=1, keepdims=True)
    normas[normas == 0] = 1
    return m / normas


# --------------------------------------------------------------------------
# Identidade do arquivo
# --------------------------------------------------------------------------


def impressao(dados: bytes) -> str:
    """SHA-256 do arquivo. É o que responde 'já processei este?'.

    Comparar por nome não serve: o mesmo livro chega como `livro.pdf`,
    `livro (1).pdf` e `Livro Final.pdf`, e livros diferentes chegam com o
    mesmo nome. E aqui o custo de errar é alto — reprocessar gasta a cota de
    vetores, que é o recurso escasso deste projeto.
    """
    return hashlib.sha256(dados).hexdigest()


def por_impressao(digital: str) -> dict[str, object] | None:
    for livro in listar():
        if livro.get("hash") == digital:
            return livro
    return None


def _slug_livre(base: str) -> str:
    """Slug que ainda não está em uso.

    Sem isto, dois livros com nome de arquivo parecido se sobrescrevem em
    silêncio — e o primeiro é perdido junto com a cota gasta para processá-lo.
    """
    if not (pasta_livros() / base).exists():
        return base
    for n in range(2, 200):
        if not (pasta_livros() / f"{base}-{n}").exists():
            return f"{base}-{n}"
    raise BibliotecaError(f"Já existem livros demais com o nome '{base}'.")


def compativel(livro: dict[str, object]) -> bool:
    """Os vetores deste livro servem para a busca de hoje?

    Trocar de modelo ou de número de dimensões invalida tudo que já foi
    gravado. E a falha é traiçoeira: com dimensão diferente a multiplicação
    estoura e o livro desaparece da busca sem avisar; com a MESMA dimensão e
    modelo diferente, a conta funciona e devolve lixo — ela citaria a página
    errada com toda a confiança.

    Livro gravado antes desta versão não tem os campos. Assumir que combina é
    correto: foi feito com estes mesmos valores.
    """
    return (
        str(livro.get("modelo", MODELO_EMBED)) == MODELO_EMBED
        and int(livro.get("dimensoes", DIMENSOES)) == DIMENSOES
    )


# --------------------------------------------------------------------------
# Adicionar um livro
# --------------------------------------------------------------------------


async def adicionar(nome_arquivo: str, dados: bytes) -> AsyncIterator[dict[str, object]]:
    """Processa um arquivo, relatando o progresso conforme avança."""
    if len(dados) > LIMITE_ARQUIVO:
        raise BibliotecaError(
            f"O arquivo tem {len(dados) // (1024 * 1024)} MB; o teto é "
            f"{LIMITE_ARQUIVO // (1024 * 1024)} MB."
        )

    titulo = Path(nome_arquivo).stem.replace("_", " ").replace("-", " ").strip()

    # Antes de qualquer trabalho: já processei este arquivo? A resposta vem do
    # conteúdo, não do nome. Vetorizar de novo gastaria a cota à toa.
    digital = impressao(dados)
    repetido = por_impressao(digital)
    if repetido:
        yield {
            "etapa": "repetido",
            "texto": f"esse arquivo já está na biblioteca como \u201c{repetido['titulo']}\u201d",
            "livro": repetido,
        }
        return

    yield {"etapa": "lendo", "texto": "extraindo o texto…"}
    paginas = extrair(nome_arquivo, dados)

    yield {"etapa": "dividindo", "texto": f"{len(paginas)} páginas, separando em trechos…"}
    trechos = dividir(paginas)
    if not trechos:
        raise BibliotecaError("O arquivo não tem texto suficiente para valer a pena.")

    total = len(trechos)
    vetores: list[list[float]] = []

    for inicio in range(0, total, LOTE):
        pedaco = trechos[inicio : inicio + LOTE]
        vetores.extend(await _embutir([str(t["texto"]) for t in pedaco], "RETRIEVAL_DOCUMENT"))
        feito = min(inicio + LOTE, total)
        yield {
            "etapa": "vetorizando",
            "texto": f"processando… {feito} de {total} trechos",
            "progresso": round(feito / total, 2),
        }

    # Gravação atômica: monta numa pasta temporária e só então move para o
    # lugar. Livro é feito de três arquivos que só servem juntos — vetores,
    # trechos e meta. Interromper no meio deixaria um livro que a busca
    # carrega e não consegue usar.
    base_livros = pasta_livros()
    base_livros.mkdir(parents=True, exist_ok=True)
    slug = _slug_livre(slugify(titulo) or "documento")
    destino = base_livros / slug
    parcial = base_livros / f".{slug}.parcial"

    if parcial.exists():
        shutil.rmtree(parcial, ignore_errors=True)
    parcial.mkdir(parents=True)

    try:
        np.save(parcial / "vetores.npy", _normalizar(np.array(vetores, dtype=np.float32)))

        with (parcial / "trechos.jsonl").open("w", encoding="utf-8") as f:
            for t in trechos:
                f.write(json.dumps(t, ensure_ascii=False) + "\n")

        meta = {
            "titulo": titulo,
            "slug": slug,
            "arquivo": nome_arquivo,
            "paginas": len(paginas),
            "trechos": total,
            "criado": date.today().isoformat(),
            # Sem estes três, trocar de modelo ou de dimensão transforma os
            # vetores antigos em lixo silencioso. Ver compativel().
            "modelo": MODELO_EMBED,
            "dimensoes": DIMENSOES,
            "hash": digital,
            "bytes": len(dados),
        }
        (parcial / "meta.json").write_text(
            json.dumps(meta, ensure_ascii=False, indent=2), encoding="utf-8"
        )

        parcial.replace(destino)
    except Exception:
        shutil.rmtree(parcial, ignore_errors=True)
        raise

    yield {"etapa": "pronto", "livro": meta}


# --------------------------------------------------------------------------
# Consultar
# --------------------------------------------------------------------------


def listar() -> list[dict[str, object]]:
    base = pasta_livros()
    if not base.exists():
        return []
    livros = []
    for pasta in sorted(base.iterdir()):
        if pasta.name.startswith("."):
            continue  # processamento interrompido, não é livro
        meta = pasta / "meta.json"
        if not meta.exists():
            continue
        try:
            livro = json.loads(meta.read_text(encoding="utf-8"))
        except json.JSONDecodeError:
            continue
        livro["compativel"] = compativel(livro)
        livros.append(livro)
    return livros


def remover(slug: str) -> bool:
    destino = pasta_livros() / slugify(slug)
    if not destino.exists() or not destino.is_dir():
        return False
    shutil.rmtree(destino, ignore_errors=True)
    return not destino.exists()


def vazia() -> bool:
    return not listar()


async def buscar(pergunta: str, quantos: int = 4) -> list[dict[str, object]]:
    """Trechos que respondem à pergunta, do mais relevante para o menos.

    Devolve lista vazia quando nada passa do limiar — ou seja, quando a
    pergunta não tem a ver com nenhum livro guardado. Isso é o normal e não
    é erro.
    """
    livros = listar()
    if not livros:
        return []

    try:
        vetor_pergunta = (await _embutir([pergunta], "RETRIEVAL_QUERY"))[0]
    except BibliotecaError:
        return []  # a biblioteca é um bônus; falhar nela não derruba a conversa

    alvo = _normalizar(np.array([vetor_pergunta], dtype=np.float32))[0]
    achados: list[dict[str, object]] = []

    for livro in livros:
        # Vetor de outro modelo daria semelhança sem significado. Pular é
        # a única saída honesta: melhor não citar do que citar errado.
        if not compativel(livro):
            continue
        pasta = pasta_livros() / str(livro["slug"])
        try:
            vetores = np.load(pasta / "vetores.npy")
            linhas = (pasta / "trechos.jsonl").read_text(encoding="utf-8").splitlines()
        except (OSError, ValueError):
            continue

        semelhancas = vetores @ alvo
        melhores = np.argsort(semelhancas)[::-1][:quantos]

        for i in melhores:
            nota = float(semelhancas[i])
            if nota < LIMIAR or i >= len(linhas):
                continue
            try:
                trecho = json.loads(linhas[i])
            except json.JSONDecodeError:
                continue
            achados.append(
                {
                    "livro": livro["titulo"],
                    "pagina": trecho.get("pagina", 0),
                    "texto": trecho.get("texto", ""),
                    "nota": round(nota, 3),
                }
            )

    achados.sort(key=lambda a: a["nota"], reverse=True)
    return achados[:quantos]


def formatar(achados: list[dict[str, object]]) -> str:
    """Vira o bloco que entra no prompt, antes da pergunta."""
    if not achados:
        return ""
    linhas = [
        "Trechos dos livros que o André guardou, escolhidos por semelhança com "
        "a pergunta abaixo. Responda com base neles quando servirem, dizendo de "
        "qual livro e página veio. Se não responderem a pergunta, ignore-os e "
        "diga que o material dele não cobre isso — não invente para preencher.",
        "",
    ]
    for a in achados:
        origem = f"{a['livro']}"
        if a["pagina"]:
            origem += f", p. {a['pagina']}"
        linhas.append(f"--- {origem} ---")
        linhas.append(str(a["texto"]).strip())
        linhas.append("")
    return "\n".join(linhas)
