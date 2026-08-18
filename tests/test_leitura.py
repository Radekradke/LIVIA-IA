"""Fase 6 — ler PDF, DOCX, XLSX, CSV, JSON e HTML.

O teste mais importante do arquivo é o de ida e volta: o que a fase 5 escreve,
a fase 6 tem que ler. Se as duas metades não se encontram, nenhuma serve.
"""

import json
import zipfile

import pytest

from livia import ferramentas, leitura, objetos
from livia.ferramentas import FerramentaError

pytestmark = pytest.mark.usefixtures("workspace")


# --------------------------------------------------------------------------
# Ida e volta
# --------------------------------------------------------------------------


def test_le_o_docx_que_ela_mesma_escreveu(workspace):
    objetos.criar_documento(
        "r.docx", "Plano de agosto", "# Metas\n\nFechar a fase 6.\n\n- testes\n- commit", "docx"
    )
    texto = ferramentas.ler("r.docx")

    assert "Plano de agosto" in texto
    assert "Fechar a fase 6." in texto
    assert "- testes" in texto, "a lista deveria voltar como lista"
    assert "# Metas" in texto, "o título deveria voltar marcado"


def test_le_o_xlsx_que_ela_mesma_escreveu(workspace):
    objetos.criar_planilha(
        "c.xlsx", "Custos", ["Item", "Valor"], [["Servidor", 0], ["Domínio", 40]], "xlsx"
    )
    texto = ferramentas.ler("c.xlsx")

    assert "aba: Custos" in texto
    assert "Item | Valor" in texto
    assert "Servidor" in texto and "40" in texto


def test_le_o_pdf_que_ela_mesma_escreveu(workspace):
    objetos.criar_documento("r.pdf", "Relatório", "Conteúdo que precisa voltar.", "pdf")
    texto = ferramentas.ler("r.pdf")

    assert "página 1" in texto
    assert "Conteúdo que precisa voltar." in texto


def test_le_o_csv_que_ela_mesma_escreveu(workspace):
    objetos.criar_planilha("l.csv", "x", ["Nome", "Cidade"], [["Ana", "Recife"]], "csv")
    texto = ferramentas.ler("l.csv")

    assert "Nome | Cidade" in texto
    assert "Ana | Recife" in texto
    assert "separador ';'" in texto


# --------------------------------------------------------------------------
# XLSX: fórmula é a armadilha
# --------------------------------------------------------------------------


def test_planilha_devolve_o_resultado_e_nao_a_formula(workspace):
    """Sem `data_only=True` o openpyxl entrega '=SOMA(A1:A2)' como se fosse o
    dado, e a Livia responderia comentando o texto da fórmula."""
    from openpyxl import load_workbook

    objetos.criar_planilha("s.xlsx", "n", ["Valor"], [[10], [20]], "xlsx")

    # Simula o que o Excel faz ao salvar: grava a fórmula E o valor calculado.
    wb = load_workbook(str(workspace / "s.xlsx"))
    wb.active["A4"] = "=SUM(A2:A3)"
    wb.save(str(workspace / "s.xlsx"))

    texto = ferramentas.ler("s.xlsx")
    assert "=SUM" not in texto, "isso é a fórmula, não o dado"


def test_planilha_sem_valor_em_cache_avisa_em_vez_de_dizer_vazia(workspace):
    """Arquivo gerado por programa não tem resultado de fórmula gravado.

    Devolver 'vazia' faria a Livia concluir que não há dado nenhum ali.
    """
    from openpyxl import Workbook

    wb = Workbook()
    wb.active["A1"] = "=1+1"
    wb.save(str(workspace / "f.xlsx"))

    texto = ferramentas.ler("f.xlsx")
    assert "fórmula" in texto.lower()
    assert "nunca foi aberto no Excel" in texto


def test_le_todas_as_abas(workspace):
    from openpyxl import Workbook

    wb = Workbook()
    wb.active.title = "Primeira"
    wb.active["A1"] = "um"
    wb.create_sheet("Segunda")["A1"] = "dois"
    wb.save(str(workspace / "m.xlsx"))

    texto = ferramentas.ler("m.xlsx")
    assert "aba: Primeira" in texto and "aba: Segunda" in texto
    assert "um" in texto and "dois" in texto


def test_planilha_grande_e_truncada(workspace):
    from openpyxl import Workbook

    wb = Workbook()
    for i in range(leitura.LIMITE_LINHAS_TABELA + 50):
        wb.active.append([i])
    wb.save(str(workspace / "g.xlsx"))

    texto = ferramentas.ler("g.xlsx")
    assert "truncada" in texto


# --------------------------------------------------------------------------
# DOCX
# --------------------------------------------------------------------------


def test_tabela_do_docx_nao_e_perdida(workspace):
    """A tabela costuma ser onde o dado está."""
    from docx import Document

    doc = Document()
    doc.add_paragraph("Segue o orçamento:")
    tabela = doc.add_table(rows=2, cols=2)
    tabela.cell(0, 0).text = "Item"
    tabela.cell(0, 1).text = "Preço"
    tabela.cell(1, 0).text = "Domínio"
    tabela.cell(1, 1).text = "40"
    doc.save(str(workspace / "t.docx"))

    texto = ferramentas.ler("t.docx")
    assert "Item | Preço" in texto
    assert "Domínio | 40" in texto


def test_docx_falso_da_erro_claro(workspace):
    (workspace / "mentira.docx").write_text("não sou um docx", encoding="utf-8")

    with pytest.raises(FerramentaError) as erro:
        ferramentas.ler("mentira.docx")
    assert "ZIP" in str(erro.value) or "corrompido" in str(erro.value)


# --------------------------------------------------------------------------
# Bomba-zip
# --------------------------------------------------------------------------


def test_bomba_zip_e_recusada_antes_de_abrir(workspace):
    """Um .docx de alguns KB pode virar gigabytes ao descompactar.

    O arquivo vem de fora — isso existe antes de qualquer má intenção do
    modelo.
    """
    caminho = workspace / "bomba.docx"
    with zipfile.ZipFile(caminho, "w", zipfile.ZIP_DEFLATED) as z:
        z.writestr("word/document.xml", b"\0" * (200 * 1024 * 1024))

    assert caminho.stat().st_size < 1024 * 1024, "o teste precisa de um arquivo pequeno"

    with pytest.raises(FerramentaError) as erro:
        ferramentas.ler("bomba.docx")
    assert "bomba-zip" in str(erro.value) or "descompactado" in str(erro.value)


def test_arquivo_gigante_e_recusado(workspace, monkeypatch):
    monkeypatch.setattr(leitura, "LIMITE_ARQUIVO", 100)
    (workspace / "g.json").write_text("x" * 500, encoding="utf-8")

    with pytest.raises(FerramentaError) as erro:
        ferramentas.ler("g.json")
    assert "grande demais" in str(erro.value)


# --------------------------------------------------------------------------
# PDF
# --------------------------------------------------------------------------


def test_pdf_escaneado_diz_a_verdade(workspace, monkeypatch):
    """Sem camada de texto, devolver vazio faria a Livia inventar o conteúdo."""
    objetos.criar_documento("scan.pdf", "T", "texto qualquer", "pdf")

    class PaginaSemTexto:
        def extract_text(self):
            return ""

    class LeitorFalso:
        is_encrypted = False
        pages = [PaginaSemTexto(), PaginaSemTexto()]

        def __init__(self, *a, **k):
            pass

    import pypdf

    monkeypatch.setattr(pypdf, "PdfReader", LeitorFalso)

    with pytest.raises(FerramentaError) as erro:
        ferramentas.ler("scan.pdf")
    msg = str(erro.value)
    assert "OCR" in msg and "escaneamento" in msg
    assert "2 página" in msg


def test_pdf_corrompido_nao_vaza_stack_trace(workspace):
    (workspace / "quebrado.pdf").write_bytes(b"%PDF-1.4\nlixo")

    with pytest.raises(FerramentaError) as erro:
        ferramentas.ler("quebrado.pdf")
    assert "Traceback" not in str(erro.value)


# --------------------------------------------------------------------------
# HTML
# --------------------------------------------------------------------------


def test_html_vira_texto_sem_o_javascript(workspace):
    """Tirar só as tags deixaria o script inteiro no meio da leitura."""
    (workspace / "p.html").write_text(
        "<html><head><title>Notícia</title>"
        "<style>body{color:red}</style></head><body>"
        "<script>var segredo='nao deveria aparecer';</script>"
        "<h1>Manchete</h1><p>Primeiro parágrafo.</p></body></html>",
        encoding="utf-8",
    )

    texto = ferramentas.ler("p.html")
    assert "nao deveria aparecer" not in texto
    assert "color:red" not in texto
    assert "Manchete" in texto and "Primeiro parágrafo." in texto
    assert "Notícia" in texto, "o título da página é contexto útil"
    assert "<" not in texto


def test_entidades_html_voltam_como_texto(workspace):
    (workspace / "e.html").write_text("<p>caf&eacute; &amp; p&atilde;o</p>", encoding="utf-8")
    assert "café & pão" in ferramentas.ler("e.html")


# --------------------------------------------------------------------------
# JSON e CSV
# --------------------------------------------------------------------------


def test_json_volta_formatado_com_resumo(workspace):
    (workspace / "d.json").write_text(
        json.dumps({"nome": "Livia", "fases": [1, 2, 3]}), encoding="utf-8"
    )
    texto = ferramentas.ler("d.json")

    assert "objeto com as chaves" in texto
    assert "nome" in texto and "fases" in texto
    assert '"nome": "Livia"' in texto, "deveria vir indentado"


def test_json_invalido_devolve_o_texto_e_aponta_o_erro(workspace):
    (workspace / "b.json").write_text('{"a": 1,}', encoding="utf-8")
    texto = ferramentas.ler("b.json")

    assert "JSON inválido" in texto
    assert "linha 1" in texto
    assert '{"a": 1,}' in texto, "o conteúdo cru ajuda o modelo a achar o erro"


def test_csv_com_virgula_tambem_funciona(workspace):
    (workspace / "us.csv").write_text("a,b,c\n1,2,3\n", encoding="utf-8")
    texto = ferramentas.ler("us.csv")

    assert "separador ','" in texto
    assert "a | b | c" in texto


def test_csv_grande_e_truncado(workspace):
    linhas = "\n".join(f"{i};x" for i in range(leitura.LIMITE_LINHAS_TABELA + 100))
    (workspace / "g.csv").write_text(linhas, encoding="utf-8")

    texto = ferramentas.ler("g.csv")
    assert "e mais 100 linhas" in texto


def test_arquivo_latin1_nao_quebra(workspace):
    """Planilha exportada de sistema antigo raramente vem em UTF-8."""
    (workspace / "velho.csv").write_bytes("nome;cidade\nJoão;Belém\n".encode("cp1252"))
    texto = ferramentas.ler("velho.csv")
    assert "João" in texto and "Belém" in texto


# --------------------------------------------------------------------------
# Confinamento e limites herdados
# --------------------------------------------------------------------------


@pytest.mark.parametrize("caminho", ["../fora.pdf", "/etc/passwd", "C:\\Windows\\win.ini"])
def test_confinamento_continua_valendo(workspace, caminho):
    with pytest.raises(FerramentaError):
        ferramentas.ler(caminho)


def test_binario_desconhecido_lista_o_que_ela_le(workspace):
    (workspace / "foto.png").write_bytes(b"\x89PNG\r\n\x1a\n\x00\x01\x02\x03")

    with pytest.raises(FerramentaError) as erro:
        ferramentas.ler("foto.png")
    msg = str(erro.value)
    assert "pdf" in msg and "xlsx" in msg, "dizer o que ela LÊ é mais útil que só recusar"


def test_leitura_longa_e_truncada(workspace, monkeypatch):
    monkeypatch.setattr(ferramentas, "LIMITE_LEITURA", 200)
    (workspace / "longo.json").write_text(json.dumps({"x": ["y"] * 500}), encoding="utf-8")

    texto = ferramentas.ler("longo.json")
    assert "truncado" in texto
    assert len(texto) < 500


def test_texto_puro_continua_como_era(workspace):
    (workspace / "nota.md").write_text("# Oi\n\nconteúdo", encoding="utf-8")
    assert ferramentas.ler("nota.md") == "# Oi\n\nconteúdo"


def test_lista_do_docx_nao_gasta_linha_em_branco_por_item(workspace):
    """Separar tudo com parágrafo vazio dobra o tamanho de um documento
    longo — e isso é janela de contexto gasta à toa."""
    objetos.criar_documento("l.docx", "T", "- um\n- dois\n- três", "docx")
    texto = ferramentas.ler("l.docx")

    assert "- um\n- dois\n- três" in texto
