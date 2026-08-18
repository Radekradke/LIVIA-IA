"""Fase 8 — biblioteca: DOCX, impressão digital, gravação atômica e o
registro de qual modelo gerou os vetores.

A cota de vetores é o recurso escasso deste projeto. Metade destes testes
existe para não gastá-la duas vezes com o mesmo arquivo; a outra metade para
não gastá-la e perder o resultado.
"""

from __future__ import annotations

import io
import json
import zipfile

import numpy as np
import pytest

from livia import biblioteca

pytestmark = pytest.mark.usefixtures("dados")


# --------------------------------------------------------------------------
# Ferramentas do teste
# --------------------------------------------------------------------------


def docx_bytes(paragrafos, tabela=None) -> bytes:
    from docx import Document

    doc = Document()
    for texto, estilo in paragrafos:
        doc.add_paragraph(texto, style=estilo) if estilo else doc.add_paragraph(texto)
    if tabela:
        t = doc.add_table(rows=len(tabela), cols=len(tabela[0]))
        for i, linha in enumerate(tabela):
            for j, celula in enumerate(linha):
                t.cell(i, j).text = celula
    buffer = io.BytesIO()
    doc.save(buffer)
    return buffer.getvalue()


@pytest.fixture
def sem_rede(monkeypatch):
    """Vetores falsos e determinísticos. O que se testa aqui é o que acontece
    em volta da chamada, não a chamada."""
    chamadas = {"n": 0, "textos": []}

    async def falso(textos, tarefa):
        chamadas["n"] += 1
        chamadas["textos"].extend(textos)
        return [
            [float((hash(t) % 1000) / 1000)] * biblioteca.DIMENSOES for t in textos
        ]

    monkeypatch.setattr(biblioteca, "_embutir", falso)
    return chamadas


async def guardar(nome, dados):
    passos = []
    async for passo in biblioteca.adicionar(nome, dados):
        passos.append(passo)
    return passos


# --------------------------------------------------------------------------
# DOCX
# --------------------------------------------------------------------------


def test_docx_agora_e_aceito():
    """A mensagem antiga mandava converter DOCX para PDF. Agora não precisa."""
    corpo = "Explicação longa o bastante para virar um trecho de verdade. " * 6
    secoes = biblioteca.extrair("apostila.docx", docx_bytes([(corpo, None)]))

    assert secoes and corpo.strip()[:40] in secoes[0][1]


def test_docx_e_dividido_por_titulo():
    """DOCX não tem página — o Word pagina na hora de imprimir. Sem uma âncora,
    a citação não aponta para lugar nenhum."""
    secoes = biblioteca.extrair(
        "manual.docx",
        docx_bytes([
            ("Introdução", "Heading 1"),
            ("Texto da introdução.", None),
            ("Instalação", "Heading 1"),
            ("Texto da instalação.", None),
        ]),
    )

    assert len(secoes) == 2
    assert [n for n, _ in secoes] == [1, 2]
    assert "introdução" in secoes[0][1].lower()
    assert "instalação" in secoes[1][1].lower()


def test_tabela_do_docx_entra_no_texto():
    secoes = biblioteca.extrair(
        "dados.docx",
        docx_bytes([("Segue a tabela:", None)], tabela=[["Item", "Preço"], ["Domínio", "40"]]),
    )
    tudo = "\n".join(t for _, t in secoes)
    assert "Item | Preço" in tudo and "Domínio | 40" in tudo


def test_docx_falso_da_erro_em_portugues():
    with pytest.raises(biblioteca.BibliotecaError) as erro:
        biblioteca.extrair("mentira.docx", b"nao sou um docx")
    assert "Traceback" not in str(erro.value)


def test_docx_vazio_avisa():
    with pytest.raises(biblioteca.BibliotecaError):
        biblioteca.extrair("vazio.docx", docx_bytes([]))


def test_bomba_zip_nao_passa_pela_biblioteca():
    """O upload vem de fora e nunca toca o disco antes da conferência."""
    buffer = io.BytesIO()
    with zipfile.ZipFile(buffer, "w", zipfile.ZIP_DEFLATED) as z:
        z.writestr("word/document.xml", b"\0" * (200 * 1024 * 1024))

    assert len(buffer.getvalue()) < 1024 * 1024

    with pytest.raises(biblioteca.BibliotecaError) as erro:
        biblioteca.extrair("bomba.docx", buffer.getvalue())
    assert "bomba-zip" in str(erro.value) or "descompactado" in str(erro.value)


def test_mensagem_de_formato_nao_manda_mais_converter_docx():
    with pytest.raises(biblioteca.BibliotecaError) as erro:
        biblioteca.extrair("livro.epub", b"x")
    msg = str(erro.value)
    assert "DOCX" in msg
    assert "converta para PDF" not in msg, "a instrução ficou obsoleta"


# --------------------------------------------------------------------------
# Impressão digital e repetição
# --------------------------------------------------------------------------


def test_impressao_depende_do_conteudo_e_nao_do_nome():
    """O mesmo livro chega como `livro.pdf`, `livro (1).pdf` e `Final.pdf`."""
    a = biblioteca.impressao(b"conteudo identico")
    b = biblioteca.impressao(b"conteudo identico")
    c = biblioteca.impressao(b"conteudo diferente")

    assert a == b and a != c
    assert len(a) == 64


@pytest.mark.asyncio
async def test_arquivo_repetido_nao_gasta_a_cota_de_novo(sem_rede):
    """Vetorizar de novo o mesmo arquivo é jogar a cota fora."""
    conteudo = ("Explicação sobre ponteiros em C, com detalhe suficiente. " * 8).encode()

    await guardar("curso-de-c.txt", conteudo)
    chamadas_primeira = sem_rede["n"]
    assert chamadas_primeira > 0

    passos = await guardar("curso-de-c-copia (1).txt", conteudo)

    assert sem_rede["n"] == chamadas_primeira, "não deveria ter chamado a API"
    assert passos[0]["etapa"] == "repetido"
    assert "já está na biblioteca" in passos[0]["texto"]
    assert len(biblioteca.listar()) == 1, "não deveria ter duplicado"


@pytest.mark.asyncio
async def test_livros_diferentes_com_nome_parecido_nao_se_sobrescrevem(sem_rede):
    """Sem slug livre, o segundo apagaria o primeiro — e a cota gasta com ele."""
    um = ("Primeiro livro, sobre ponteiros e alocação de memória. " * 8).encode()
    dois = ("Segundo livro, sobre estruturas de dados e árvores. " * 8).encode()

    await guardar("Curso de C.txt", um)
    await guardar("curso-de-c.txt", dois)

    livros = biblioteca.listar()
    assert len(livros) == 2, "um sobrescreveu o outro"
    assert len({str(l["slug"]) for l in livros}) == 2
    assert len({str(l["hash"]) for l in livros}) == 2


@pytest.mark.asyncio
async def test_arquivo_grande_e_recusado_antes_de_processar(sem_rede):
    grande = b"x" * (biblioteca.LIMITE_ARQUIVO + 1)

    with pytest.raises(biblioteca.BibliotecaError) as erro:
        await guardar("enorme.txt", grande)

    assert "MB" in str(erro.value)
    assert sem_rede["n"] == 0, "nem deveria ter começado"


# --------------------------------------------------------------------------
# O que meta.json passou a guardar
# --------------------------------------------------------------------------


@pytest.mark.asyncio
async def test_meta_registra_modelo_dimensoes_e_hash(sem_rede):
    conteudo = ("Texto suficiente para gerar pelo menos um trecho válido. " * 8).encode()
    passos = await guardar("apostila.txt", conteudo)
    meta = passos[-1]["livro"]

    assert meta["modelo"] == biblioteca.MODELO_EMBED
    assert meta["dimensoes"] == biblioteca.DIMENSOES
    assert meta["hash"] == biblioteca.impressao(conteudo)
    assert meta["bytes"] == len(conteudo)


@pytest.mark.asyncio
async def test_troca_de_modelo_marca_o_livro_como_incompativel(sem_rede, monkeypatch):
    """Aqui está a falha traiçoeira: com a MESMA dimensão e modelo diferente, a
    multiplicação funciona e devolve lixo. Ela citaria a página errada com toda
    a confiança."""
    await guardar("livro.txt", ("Conteúdo com tamanho suficiente para um trecho. " * 8).encode())
    assert biblioteca.listar()[0]["compativel"] is True

    monkeypatch.setattr(biblioteca, "MODELO_EMBED", "outro-modelo-de-vetores")

    livro = biblioteca.listar()[0]
    assert livro["compativel"] is False
    assert biblioteca.compativel(livro) is False


@pytest.mark.asyncio
async def test_troca_de_dimensao_tambem_invalida(sem_rede, monkeypatch):
    await guardar("livro.txt", ("Conteúdo com tamanho suficiente para um trecho. " * 8).encode())
    monkeypatch.setattr(biblioteca, "DIMENSOES", 1536)
    assert biblioteca.listar()[0]["compativel"] is False


@pytest.mark.asyncio
async def test_busca_ignora_livro_incompativel(sem_rede, monkeypatch):
    """Melhor não citar do que citar errado.

    O teste prova as duas metades. Sem a primeira, ele passaria também se
    `buscar()` nunca devolvesse nada — e não seria prova de coisa alguma.
    """
    await guardar("livro.txt", ("Ponteiros, alocação e liberação de memória em C. " * 8).encode())

    # Limiar no chão: qualquer semelhança serve, então o trecho TEM que vir.
    monkeypatch.setattr(biblioteca, "LIMIAR", -1.0)
    achados = await biblioteca.buscar("ponteiros")
    assert achados, "o livro compatível deveria ter sido encontrado"

    # Mesmo livro, mesmo limiar, só o modelo mudou.
    monkeypatch.setattr(biblioteca, "MODELO_EMBED", "outro-modelo")
    assert await biblioteca.buscar("ponteiros") == [], "citou de um livro invalidado"


def test_livro_antigo_sem_os_campos_continua_valendo():
    """Livro gravado antes desta versão não tem `modelo` nem `dimensoes`.
    Assumir que combina é correto: foi feito com estes mesmos valores."""
    assert biblioteca.compativel({"titulo": "Antigo", "slug": "antigo"}) is True


# --------------------------------------------------------------------------
# Gravação atômica
# --------------------------------------------------------------------------


@pytest.mark.asyncio
async def test_falha_ao_gravar_nao_deixa_livro_pela_metade(sem_rede, monkeypatch):
    """Livro é três arquivos que só servem juntos. Um livro incompleto seria
    carregado pela busca e não daria para usar."""
    def explode(*a, **k):
        raise OSError("disco cheio")

    monkeypatch.setattr(np, "save", explode)

    with pytest.raises(OSError):
        await guardar("livro.txt", ("Texto longo o suficiente para um trecho. " * 8).encode())

    assert biblioteca.listar() == []
    sobras = list(biblioteca.pasta_livros().iterdir()) if biblioteca.pasta_livros().exists() else []
    assert sobras == [], f"sobrou lixo: {sobras}"


@pytest.mark.asyncio
async def test_pasta_parcial_de_execucao_anterior_e_ignorada(sem_rede):
    """Interrupção no meio deixa `.slug.parcial`. Aquilo não é livro."""
    biblioteca.pasta_livros().mkdir(parents=True, exist_ok=True)
    restos = biblioteca.pasta_livros() / ".livro.parcial"
    restos.mkdir()
    (restos / "meta.json").write_text('{"titulo": "Metade", "slug": "livro"}', encoding="utf-8")

    assert biblioteca.listar() == [], "pasta parcial apareceu como livro"


@pytest.mark.asyncio
async def test_o_livro_gravado_tem_as_tres_partes(sem_rede):
    conteudo = ("Conteúdo bastante para gerar trechos de verdade aqui. " * 10).encode()
    passos = await guardar("apostila.txt", conteudo)
    pasta = biblioteca.pasta_livros() / str(passos[-1]["livro"]["slug"])

    assert (pasta / "vetores.npy").exists()
    assert (pasta / "trechos.jsonl").exists()
    assert (pasta / "meta.json").exists()

    vetores = np.load(pasta / "vetores.npy")
    linhas = (pasta / "trechos.jsonl").read_text(encoding="utf-8").splitlines()
    assert vetores.shape == (len(linhas), biblioteca.DIMENSOES)
    assert json.loads(linhas[0])["texto"]


@pytest.mark.asyncio
async def test_remover_apaga_o_livro_inteiro(sem_rede):
    passos = await guardar("apostila.txt", ("Texto grande o bastante aqui. " * 10).encode())
    slug = str(passos[-1]["livro"]["slug"])

    assert biblioteca.remover(slug) is True
    assert not (biblioteca.pasta_livros() / slug).exists()
    assert biblioteca.listar() == []


@pytest.mark.asyncio
async def test_remover_aguenta_subpasta(sem_rede):
    """A versão anterior usava unlink() em tudo e quebraria numa subpasta."""
    passos = await guardar("apostila.txt", ("Texto grande o bastante aqui. " * 10).encode())
    slug = str(passos[-1]["livro"]["slug"])
    (biblioteca.pasta_livros() / slug / "extra").mkdir()
    (biblioteca.pasta_livros() / slug / "extra" / "nota.txt").write_text("x", encoding="utf-8")

    assert biblioteca.remover(slug) is True


def test_remover_o_que_nao_existe_devolve_falso():
    assert biblioteca.remover("nunca-existiu") is False
