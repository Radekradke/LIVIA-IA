"""Fase 5 — gerar documentos e planilhas de verdade.

O teste não confere só se o arquivo apareceu: abre cada um com a biblioteca
que o formato exige. Um DOCX corrompido tem tamanho maior que zero.
"""

import csv
import zipfile

import pytest

from livia import ferramentas, objetos
from livia.ferramentas import FerramentaError

pytestmark = pytest.mark.usefixtures("workspace")


# --------------------------------------------------------------------------
# Cada formato precisa abrir no programa dele
# --------------------------------------------------------------------------


def test_docx_abre_no_word(workspace):
    from docx import Document

    objetos.criar_documento(
        "relatorio.docx",
        "Relatório de agosto",
        "# Resumo\n\nO mês fechou no azul.\n\n- corte de custos\n- venda nova",
        "docx",
    )

    doc = Document(str(workspace / "relatorio.docx"))
    textos = [p.text for p in doc.paragraphs]
    assert "Relatório de agosto" in textos
    assert "O mês fechou no azul." in textos
    assert "corte de custos" in textos

    estilos = [p.style.name for p in doc.paragraphs]
    assert "List Bullet" in estilos, "a lista deveria virar lista, não parágrafo"


def test_pdf_e_um_pdf_de_verdade(workspace):
    objetos.criar_documento("doc.pdf", "Teste", "Primeiro parágrafo.\n\n- item", "pdf")

    dados = (workspace / "doc.pdf").read_bytes()
    assert dados.startswith(b"%PDF-"), "sem o cabeçalho não é PDF"
    assert b"%%EOF" in dados[-2048:], "PDF truncado"

    # Se abre no leitor, abre para o André.
    from pypdf import PdfReader

    leitor = PdfReader(str(workspace / "doc.pdf"))
    assert "Primeiro parágrafo." in leitor.pages[0].extract_text()


def test_xlsx_abre_no_excel(workspace):
    from openpyxl import load_workbook

    objetos.criar_planilha(
        "custos.xlsx",
        "Agosto",
        ["Item", "Valor"],
        [["Servidor", 0], ["Domínio", 40]],
        "xlsx",
    )

    wb = load_workbook(str(workspace / "custos.xlsx"))
    aba = wb.active
    assert aba.title == "Agosto"
    assert [c.value for c in aba[1]] == ["Item", "Valor"]
    assert aba["A2"].value == "Servidor"
    assert aba.freeze_panes == "A2", "cabeçalho deveria ficar fixo ao rolar"


def test_csv_sai_legivel_no_excel_brasileiro(workspace):
    objetos.criar_planilha("lista.csv", "x", ["Nome", "Cidade"], [["Ana", "Recife"]], "csv")

    bruto = (workspace / "lista.csv").read_bytes()
    assert bruto.startswith(b"\xef\xbb\xbf"), "sem BOM o Excel erra a acentuação"

    with (workspace / "lista.csv").open(encoding="utf-8-sig", newline="") as f:
        linhas = list(csv.reader(f, delimiter=";"))
    assert linhas == [["Nome", "Cidade"], ["Ana", "Recife"]]


def test_html_escapa_o_conteudo(workspace):
    objetos.criar_documento(
        "pagina.html", "Teste", "Cuidado com <script>alert(1)</script> solto.", "html"
    )

    texto = (workspace / "pagina.html").read_text(encoding="utf-8")
    assert "<script>alert(1)</script>" not in texto, "isso executaria ao abrir"
    assert "&lt;script&gt;" in texto


def test_markdown_vira_estrutura_no_html(workspace):
    objetos.criar_documento(
        "p.html", "T", "# Cabeça\n\ntexto\n\n- um\n- dois\n\n1. primeiro", "html"
    )
    texto = (workspace / "p.html").read_text(encoding="utf-8")

    assert "<h2>Cabeça</h2>" in texto
    assert texto.count("<li>") == 3
    assert "<ul>" in texto and "<ol>" in texto, "marcador e número são listas diferentes"
    assert "<p>texto</p>" in texto


# --------------------------------------------------------------------------
# Injeção de fórmula
# --------------------------------------------------------------------------


@pytest.mark.parametrize(
    "ataque",
    [
        "=cmd|'/c calc'!A1",
        "+1+1",
        "-1+1",
        "@SUM(A1:A9)",
        "=HYPERLINK(\"http://mal.example\",\"clique\")",
    ],
)
def test_formula_nao_entra_na_planilha(workspace, ataque):
    """Célula começando com = vira execução quando alguém abre no Excel.

    O conteúdo vem do modelo, que por sua vez leu a internet. Confiar nisso
    seria irresponsável.
    """
    from openpyxl import load_workbook

    objetos.criar_planilha("x.xlsx", "a", ["Campo"], [[ataque]], "xlsx")
    valor = load_workbook(str(workspace / "x.xlsx")).active["A2"].value

    assert isinstance(valor, str)
    assert not valor.startswith(("=", "+", "-", "@")), f"{valor!r} ainda é fórmula"
    assert ataque in valor, "o texto original deveria continuar legível"


def test_formula_tambem_barrada_no_csv(workspace):
    objetos.criar_planilha("x.csv", "a", ["Campo"], [["=1+1"]], "csv")
    texto = (workspace / "x.csv").read_text(encoding="utf-8-sig")
    assert "'=1+1" in texto


def test_cabecalho_tambem_e_neutralizado(workspace):
    from openpyxl import load_workbook

    objetos.criar_planilha("x.xlsx", "a", ["=EVIL()"], [["ok"]], "xlsx")
    assert not load_workbook(str(workspace / "x.xlsx")).active["A1"].value.startswith("=")


def test_numero_continua_numero(workspace):
    """Neutralizar não pode transformar toda planilha em texto."""
    from openpyxl import load_workbook

    objetos.criar_planilha("n.xlsx", "a", ["Valor"], [[42]], "xlsx")
    assert load_workbook(str(workspace / "n.xlsx")).active["A2"].value == 42


# --------------------------------------------------------------------------
# Confinamento
# --------------------------------------------------------------------------


@pytest.mark.parametrize(
    "caminho",
    ["../fora.docx", "../../etc/nota.txt", "/etc/nota.txt", "C:\\Windows\\nota.txt"],
)
def test_nao_escreve_fora_da_pasta(workspace, caminho):
    with pytest.raises(FerramentaError):
        objetos.criar_documento(caminho, "T", "conteúdo", caminho.rsplit(".", 1)[1])


def test_nada_vaza_para_fora(workspace):
    antes = set(p.name for p in workspace.parent.iterdir())
    for caminho in ("../fuga.txt", "../../fuga.csv"):
        try:
            objetos.criar_documento(caminho, "T", "x", "txt")
        except FerramentaError:
            pass
    assert set(p.name for p in workspace.parent.iterdir()) == antes


def test_cria_as_pastas_do_caminho(workspace):
    objetos.criar_documento("relatorios/2026/agosto.md", "Agosto", "texto", "md")
    assert (workspace / "relatorios" / "2026" / "agosto.md").exists()


# --------------------------------------------------------------------------
# Extensão x formato
# --------------------------------------------------------------------------


def test_extensao_precisa_combinar_com_o_formato(workspace):
    with pytest.raises(FerramentaError) as erro:
        objetos.criar_documento("nota.txt", "T", "x", "docx")
    assert "txt" in str(erro.value) and "docx" in str(erro.value)
    assert not (workspace / "nota.txt").exists()


def test_formato_desconhecido_e_recusado(workspace):
    with pytest.raises(FerramentaError) as erro:
        objetos.criar_documento("v.exe", "T", "x", "exe")
    assert "exe" in str(erro.value)


def test_planilha_nao_aceita_formato_de_documento(workspace):
    with pytest.raises(FerramentaError):
        objetos.criar_planilha("t.docx", "a", ["A"], [["1"]], "docx")


# --------------------------------------------------------------------------
# Escrita atômica
# --------------------------------------------------------------------------


def test_falha_no_meio_nao_destroi_o_que_existia(workspace, monkeypatch):
    """Regenerar um relatório não pode custar a versão anterior."""
    alvo = workspace / "importante.docx"
    objetos.criar_documento("importante.docx", "Original", "primeira versão", "docx")
    original = alvo.read_bytes()

    def explode(destino, titulo, conteudo):
        destino.write_bytes(b"lixo pela metade")
        raise RuntimeError("acabou a memória")

    monkeypatch.setattr(objetos, "_escrever_docx", explode)

    with pytest.raises(FerramentaError):
        objetos.criar_documento("importante.docx", "Nova", "segunda versão", "docx")

    assert alvo.read_bytes() == original, "a versão boa foi perdida"


def test_falha_nao_deixa_arquivo_pela_metade(workspace, monkeypatch):
    def explode(destino, titulo, conteudo):
        destino.write_bytes(b"lixo")
        raise RuntimeError("falhou")

    monkeypatch.setattr(objetos, "_escrever_docx", explode)

    with pytest.raises(FerramentaError):
        objetos.criar_documento("novo.docx", "T", "x", "docx")

    assert list(workspace.iterdir()) == [], f"sobrou lixo: {list(workspace.iterdir())}"


def test_sobrescrever_guarda_a_versao_anterior(workspace):
    objetos.criar_documento("nota.md", "V1", "primeira", "md")
    objetos.criar_documento("nota.md", "V2", "segunda", "md")

    copias = list(workspace.glob("nota.antes-*.md"))
    assert len(copias) == 1
    assert "primeira" in copias[0].read_text(encoding="utf-8")
    assert "segunda" in (workspace / "nota.md").read_text(encoding="utf-8")


def test_a_copia_mantem_a_extensao(workspace):
    """`nota.antes-x.md` ainda abre com dois cliques; `nota.md.antes-x` não."""
    objetos.criar_documento("nota.docx", "V1", "a", "docx")
    objetos.criar_documento("nota.docx", "V2", "b", "docx")

    copia = next(workspace.glob("nota.antes-*"))
    assert copia.suffix == ".docx"
    assert zipfile.is_zipfile(copia), "a cópia deveria continuar sendo um DOCX válido"


# --------------------------------------------------------------------------
# Limites e formato dos dados
# --------------------------------------------------------------------------


def test_conteudo_gigante_e_recusado(workspace):
    with pytest.raises(FerramentaError) as erro:
        objetos.criar_documento("g.md", "T", "x" * (objetos.LIMITE_CONTEUDO + 1), "md")
    assert str(objetos.LIMITE_CONTEUDO) in str(erro.value)


def test_planilha_precisa_de_coluna(workspace):
    with pytest.raises(FerramentaError):
        objetos.criar_planilha("v.xlsx", "a", [], [], "xlsx")


def test_linhas_demais_sao_recusadas(workspace):
    linhas = [["x"]] * (objetos.LIMITE_LINHAS + 1)
    with pytest.raises(FerramentaError):
        objetos.criar_planilha("g.xlsx", "a", ["A"], linhas, "xlsx")


def test_linha_maior_que_o_cabecalho_e_recusada(workspace):
    with pytest.raises(FerramentaError) as erro:
        objetos.criar_planilha("d.xlsx", "a", ["A", "B"], [["1", "2", "3"]], "xlsx")
    assert "linha 1" in str(erro.value).lower()


def test_linha_curta_e_completada(workspace):
    from openpyxl import load_workbook

    objetos.criar_planilha("c.xlsx", "a", ["A", "B", "C"], [["1"]], "xlsx")
    aba = load_workbook(str(workspace / "c.xlsx")).active
    assert aba.max_column == 3
    assert aba["C2"].value is None


def test_nome_de_aba_invalido_e_ajustado(workspace):
    from openpyxl import load_workbook

    objetos.criar_planilha("a.xlsx", "Custos/2026: [rascunho] " + "x" * 40, ["A"], [["1"]], "xlsx")
    titulo = load_workbook(str(workspace / "a.xlsx")).active.title

    assert len(titulo) <= 31
    assert not set(titulo) & set("/\\*?[]:"), "o Excel recusaria abrir"


def test_planilha_vazia_de_linhas_ainda_gera_cabecalho(workspace):
    from openpyxl import load_workbook

    objetos.criar_planilha("so-cabecalho.xlsx", "a", ["Nome"], [], "xlsx")
    assert load_workbook(str(workspace / "so-cabecalho.xlsx")).active["A1"].value == "Nome"


def test_documento_vazio_nao_quebra_o_pdf(workspace):
    objetos.criar_documento("vazio.pdf", "", "", "pdf")
    assert (workspace / "vazio.pdf").read_bytes().startswith(b"%PDF-")


# --------------------------------------------------------------------------
# Integração com a cadeia de ferramentas
# --------------------------------------------------------------------------


def test_o_modelo_enxerga_as_ferramentas_novas():
    nomes = [f["name"] for f in ferramentas.catalogo()]
    assert "criar_documento" in nomes and "criar_planilha" in nomes
    assert "ler_arquivo" in nomes, "as antigas não podem sumir"


def test_executar_despacha_para_objetos(workspace):
    texto, ok = ferramentas.executar(
        "criar_documento",
        {"caminho": "via-ferramenta.md", "titulo": "T", "conteudo": "oi", "formato": "md"},
    )
    assert ok and (workspace / "via-ferramenta.md").exists()
    assert "via-ferramenta.md" in texto


def test_erro_volta_como_texto_e_nao_derruba(workspace):
    """O modelo lê a mensagem e costuma se corrigir sozinho."""
    texto, ok = ferramentas.executar(
        "criar_planilha",
        {"caminho": "x.txt", "colunas": ["A"], "linhas": [["1"]], "formato": "xlsx"},
    )
    assert not ok
    assert "txt" in texto and "xlsx" in texto


def test_resumir_descreve_as_novas(workspace):
    assert "planilha" in ferramentas.resumir(
        "criar_planilha", {"caminho": "a.xlsx", "linhas": [[1], [2]]}
    )
    assert "documento" in ferramentas.resumir("criar_documento", {"caminho": "a.md"})
    assert "ler_arquivo" not in ferramentas.resumir("criar_documento", {"caminho": "a.md"})


def test_a_resposta_diz_onde_o_arquivo_ficou(workspace):
    texto = objetos.criar_documento("pasta/sub/nota.md", "T", "x", "md")
    assert "pasta/sub/nota.md" in texto
    assert str(workspace) not in texto, "caminho absoluto expõe a máquina"
